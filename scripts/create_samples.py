"""
Create sample test documents for development.
"""

import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
import fitz  # PyMuPDF

project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from src.config import RAW_DATA_DIR


def create_sample_pdf() -> None:
    """Create a sample PDF for testing."""
    RAW_DATA_DIR.mkdir(parents=True, exist_ok=True)
    pdf_path = RAW_DATA_DIR / "Sample_Guide.pdf"

    doc = fitz.open()
    page = doc.new_page()

    text = """Sample API Guide

1. Introduction

This is a sample API documentation for testing the document loader.
The PDF contains multiple sections and paragraphs.

2. Getting Started

To get started with our API:
1. Sign up for an account
2. Generate an API key
3. Make your first request

3. Authentication

All API requests must include your API key in the Authorization header.
We support Bearer token authentication.

4. Endpoints

The following endpoints are available:
- GET /api/users - List all users
- POST /api/users - Create a new user
- GET /api/users/{id} - Get a specific user
- PUT /api/users/{id} - Update a user
- DELETE /api/users/{id} - Delete a user

5. Error Handling

Errors are returned as JSON with appropriate HTTP status codes.
- 400: Bad Request
- 401: Unauthorized
- 404: Not Found
- 500: Server Error

6. Rate Limiting

API requests are rate-limited to 1000 requests per hour per API key.

7. Support

For support, contact: support@example.com"""

    page.insert_text((50, 50), text, fontsize=11, color=(0, 0, 0))
    doc.save(pdf_path)
    doc.close()

    print(f"✅ Created: {pdf_path}")


def create_sample_docx() -> None:
    """Create a sample DOCX file for testing."""
    RAW_DATA_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = RAW_DATA_DIR / "Sample_Vendor_Onboarding.docx"

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("Vendor Onboarding Procedure")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.color.rgb = RGBColor(0, 0, 0)

    # Content
    doc.add_paragraph("", style='List Number')

    doc.add_heading("Introduction", level=2)
    doc.add_paragraph(
        "This document outlines the complete vendor onboarding process. "
        "Please follow each step carefully to ensure successful integration."
    )

    doc.add_heading("Step 1: Initial Contact", level=2)
    doc.add_paragraph(
        "Contact our vendor management team at vendors@example.com"
    )
    doc.add_paragraph(
        "Provide company information and initial requirements",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Schedule an initial discovery call",
        style='List Bullet'
    )

    doc.add_heading("Step 2: Documentation", level=2)
    doc.add_paragraph(
        "Submit required documentation:",
        style='List Bullet'
    )
    doc.add_paragraph("Business Registration", style='List Bullet 2')
    doc.add_paragraph("Tax ID", style='List Bullet 2')
    doc.add_paragraph("W-9 Form", style='List Bullet 2')
    doc.add_paragraph("Insurance Certificate", style='List Bullet 2')

    doc.add_heading("Step 3: Legal Agreement", level=2)
    doc.add_paragraph(
        "Our legal team will prepare a vendor agreement."
    )
    doc.add_paragraph("Both parties review and sign the agreement")

    doc.add_heading("Step 4: System Setup", level=2)
    doc.add_paragraph(
        "We will set up your vendor portal access",
        style='List Bullet'
    )
    doc.add_paragraph(
        "Training on portal features and processes",
        style='List Bullet'
    )

    doc.add_heading("Step 5: Go Live", level=2)
    doc.add_paragraph(
        "Begin submitting invoices and performing transactions"
    )

    doc.add_heading("Support", level=2)
    doc.add_paragraph(
        "For onboarding support: onboarding@example.com"
    )
    doc.add_paragraph(
        "For technical support: support@example.com"
    )

    doc.save(docx_path)
    print(f"✅ Created: {docx_path}")


if __name__ == "__main__":
    print("Creating sample test documents...\n")
    create_sample_pdf()
    create_sample_docx()
    print("\n✅ Sample documents created successfully!")
