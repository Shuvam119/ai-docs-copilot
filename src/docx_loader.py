"""
DOCX Document Loader

Extracts text from Word documents (.docx) using python-docx.
"""

from docx import Document
from pathlib import Path
from typing import Dict


def load_docx(file_path: str) -> Dict:
    """
    Load a DOCX file and extract all text.

    Args:
        file_path: Path to the DOCX file

    Returns:
        Dictionary with title, text, and metadata
    """
    file_path = Path(file_path)

    if not file_path.exists():
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    if not file_path.suffix.lower() == '.docx':
        raise ValueError(f"File is not a DOCX: {file_path}")

    try:
        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n".join(paragraphs)

        return {
            "title": file_path.name,
            "text": full_text,
            "metadata": {
                "source": str(file_path),
                "type": "docx",
                "filename": file_path.name,
                "paragraphs": len(paragraphs),
                "empty_text": not full_text.strip(),
            }
        }

    except Exception as e:
        raise RuntimeError(f"Error reading DOCX {file_path}: {str(e)}")
